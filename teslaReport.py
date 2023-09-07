from pathlib import Path
from tkinter import *
import urllib.request
import csv
import os
import hashlib
import json
from docx import Document
import shutil
from datetime import date
from datetime import datetime
from docx.shared import RGBColor
from re import findall
from subprocess import Popen, PIPE
from multiprocessing import Process, Queue
import ctypes as ct
import threading


print()
print("   ______          __      ____                        __ ")
print("  /_  __/__  _____/ /___ _/ __ \___  ____  ____  _____/ /_")
print("   / / / _ \/ ___/ / __ `/ /_/ / _ \/ __ \/ __ \/ ___/ __/")
print("  / / /  __(__  ) / /_/ / _, _/  __/ /_/ / /_/ / /  / /_  ")
print(" /_/  \___/____/_/\__,_/_/ |_|\___/ .___/\____/_/   \__/  ")
print("                                 /_/                      ")
print(" By ST-R Motors")
print()

#try:
#    print(" Testmode Enabled.")
#    #urllib.request.urlretrieve("http://192.168.90.100:4035/get_data_values?format=csv", "data_values")
#    #urllib.request.urlretrieve("http://192.168.90.100:7654/diag_vitals?format=json", "diag_vitals")
#except:
#    print("[ERROR] Grabbing data values from vehicle failed. Please ensure that the SecEth is unlocked and cable is plugged.")
#    print()
#    quit()

effectiveCapacity = 0
usableCapacity = 0
acChargeCount = 0
dcChargeCount = 0
kwDischargeCount = 0
odometer= 0
carType = ""
trimName = ""
vin = ""

global brickMax
global brickMin
global brickDelta
global brickError

brickMax = 0
brickMin = 0
brickDelta = 0
brickError = 0

isConnected = False

def cprint(conText):
    canvas.itemconfig(bottomConsole, text=conText)

def tryConnection():
    global isConnected
    if isConnected:
        isConnected = False
        connectionButton.config(text="")
        connectionButton.config(image=notConnectedBG)
        cprint("Bağlantı kesildi.")
        if activePage == "general":
            goBatteryPage()
            goGeneralPage()
        elif activePage == "battery":
            goGeneralPage()
            goBatteryPage()
        elif activePage == "warning":
            goGeneralPage()
            goWarningPage()
        else:
            return
    else:
        data = ""
        cprint("MCU bağlantısı test ediliyor...")
        output = Popen(f"ping 8.8.8.8 -n 2", stdout=PIPE, encoding="utf-8")

        for line in output.stdout:
            data = data + line
            ping_test = findall("TTL", data)

        if ping_test:
            cprint("MCU tespit edildi. Veri çekilmeye çalışılıyor...")
        
            #print(" Test modu aktif.")
            urllib.request.urlretrieve("http://192.168.90.100:4035/get_data_values?format=csv", "data_values")
            urllib.request.urlretrieve("http://192.168.90.100:7654/diag_vitals?format=json", "diag_vitals")
            isConnected = True
            
            with open('data_values', encoding="utf8") as file:
                content = file.readlines()
            
            header = content[:1]
            
            global rows
            rows = content[1:]

            jsonFile = open('diag_vitals')
            
            global diagVitals
            diagVitals = json.load(jsonFile)
            
            global vin
            vin = diagVitals['vin']

            if os.path.isfile("alerts_cache"):
                os.remove("alerts_cache")
            if os.path.isfile("report_cache"):
                os.remove("report_cache")
            
            shutil.copy('basereport', 'report_cache')
            writeToReport("$datetime",str(date.today()))

            for row in rows:
                row = row.split(",")
                if row[0]=="VAPI_carType":
                    carType = row[1]
                    cprint("Araç modeli çözümleniyor...")
                    writeToReport("$vehmodel",carType)

            for row in rows:
                row = row.split(",")
                if row[0]=="GUI_odometer":
                    odometer = str(float(row[1])*1.609344)
                    cprint("Kilometre sayacı okunuyor...")
                    writeToReport("$odometer",odometer.split(".")[0]+" Km")
            
            global effectiveCapacity
            global usableCapacity
            
            for row in rows:
                row = row.split(",")
                if row[0]=="VAPI_trim":
                    trimName = row[1]
                    cprint("Araç trimi çözümleniyor...")
                    writeToReport("$vehtrim",trimName)
                    if trimName == "100D\n" or trimName == "P100D\n":
                        effectiveCapacity = 102.4
                        usableCapacity = 98.4
                    elif trimName == "90D\n" or trimName == "P90D\n":
                        effectiveCapacity = 85.8
                        usableCapacity = 81.8
                    elif trimName == "85D\n" or trimName == "P85D\n" or trimName == "85\n" or trimName == "P85\n" or trimName == "P85+\n":
                        effectiveCapacity = 81.5
                        usableCapacity = 77.5
                    elif trimName == "75D\n" or trimName == "75\n":
                        effectiveCapacity = 75
                        usableCapacity = 72.6
                    elif trimName == "70D\n" or trimName == "70\n":
                        effectiveCapacity = 75
                        usableCapacity = 65.9
                    elif trimName == "60D\n" or trimName == "60\n":
                        effectiveCapacity = 75
                        usableCapacity = 62.4
                    else:
                        cprint("Hatalı veri! Tekrar kurulum yapın.")
                        quit()
            
            cprint("Şase numarası çözümleniyor...")
            mcuVer = str(diagVitals['mcu_ver']).split(" ")[0]
            cprint(" MCU Yazılım Versiyonu: " + mcuVer)
            bdayUTC = int(diagVitals['bdayUTC'])
            bday = datetime.fromtimestamp(bdayUTC)
            bdayStr = bday.strftime('%Y-%m-%d')
            cprint(" Birthday: "+str(bday))
            writeToReport("$vin",vin)
            writeToReport("$bday",str(bday))
            writeToReport("$mcuver",mcuVer)

            generalWarrYear = add_years(bday,4)
            generalWarrKm = 50000*1.609344

            safetyWarrYear = add_years(bday,5)
            safetyWarrKm = 60000*1.609344

            batteryWarrYear = add_years(bday,8)
            batteryWarrKm = 150000*1.609344

            if datetime.now()<generalWarrYear and int(float(odometer))<int(float(generalWarrKm)):
                writeToReport("$generalwarranty","Mevcut")
                writeToReport("$generalexpiry",str(generalWarrYear) + " " + str(int(float(generalWarrKm))) + " Km")
                #print(" General Warranty: True")
            else:
                writeToReport("$generalwarranty","Mevcut Değil")
                writeToReport("$generalexpiry",str(generalWarrYear) + " " + str(int(float(generalWarrKm))) + " Km")
                #print(" General Warranty: False")

            if datetime.now()<safetyWarrYear and int(float(odometer))<int(float(safetyWarrKm)):
                writeToReport("$safetywarranty","Mevcut")
                writeToReport("$safetyexpiry",str(safetyWarrYear) + " " + str(int(float(safetyWarrKm))) + " Km")
                #print(" Safety Warranty: True")
            else:
                writeToReport("$safetywarranty","Mevcut Değil")
                writeToReport("$safetyexpiry",str(safetyWarrYear) + " " + str(int(float(safetyWarrKm))) + " Km")
                #print(" Safety Warranty: False")

            if datetime.now()<batteryWarrYear and int(float(odometer))<int(float(batteryWarrKm)):
                writeToReport("$batterywarranty","Mevcut")
                writeToReport("$batteryexpiry",str(batteryWarrYear) + " " + str(int(float(batteryWarrKm))) + " Km")
                #print(" Battery Warranty: True")
            else:
                writeToReport("$batterywarranty","Mevcut Değil")
                writeToReport("$batteryexpiry",str(batteryWarrYear) + " " + str(int(float(batteryWarrKm))) + " Km")
                #print(" Battery Warranty: False")
            
            cprint("Bağlantı kuruldu. VIN: "+vin)
            connectionButton.config(image=connectedBG)
            connectionButton.config(text="       "+vin, font= ("Inter Semibold", 14 * -1), fg="white", padx=20)
            if activePage == "general":
                goBatteryPage()
                goGeneralPage()
            elif activePage == "battery":
                goGeneralPage()
                goBatteryPage()
            elif activePage == "warning":
                goGeneralPage()
                goWarningPage()
            else:
                return
        else:
            cprint("Bağlantı kurulamadı. Lütfen SecETH kilidini açın ve kabloyu bağlayın.")
        
def add_years(start_date, years):
    try:
        return start_date.replace(year=start_date.year + years)
    except ValueError:
        return start_date.replace(year=start_date.year + years, day=28)

def writeToReport(reportVar, replaceVar):

    document = Document("report_cache")

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if reportVar in paragraph.text:
                        inline = paragraph.runs
                        paragraph.text = paragraph.text.replace(reportVar, replaceVar.strip())
    
    document.save("report_cache")

def generalInfo():
    global rows
    print()
    for row in rows:
        row = row.split(",")
        if row[0]=="VAPI_chargerType":
            print(" AC Charger Type: "+row[1])
        if row[0]=="VAPI_airSuspension":
            print(" Air Suspension: "+row[1])
        if row[0]=="VAPI_frontFogLights":
            print(" Front Fog Lights: "+row[1])
        if row[0]=="VAPI_rearFogLights":
            print(" Rear Fog Lights: "+row[1])
        if row[0]=="VAPI_hasHomelink":
            print(" Homelink: "+row[1])
        if row[0]=="VAPI_hasSunroof":
            print(" Sunroof: "+row[1])
        if row[0]=="VAPI_hasPowerLiftgate":
            print(" Power Liftgate: "+row[1])
        if row[0]=="FEATURE_blindspotWarningEnabled":
            print(" Blindspot Warning: "+row[1])
        if row[0]=="VAPI_hasMemorySeats":
            print(" Memory Seats: "+row[1])
        if row[0]=="VAPI_hasMemoryMirrors":
            print(" Memory Mirrors: "+row[1])
        if row[0]=="SYS_IC_cpuHardware":
            print(" CPU Hardware: "+row[1])
        if row[0]=="VAPI_hasSeatHeaters":
            print(" Front Seat Heater: "+row[1])
        if row[0]=="VAPI_rearSeatHeaters":
            print(" Rear Seat Heater: "+row[1])
        if row[0]=="VAPI_steeringWheelHeater":
            print(" Steering Wheel Heater: "+row[1])
        if row[0]=="VAPI_fourWheelDrive":
            print(" Four Wheel Drive: "+row[1])
        if row[0]=="VAPI_wheelType":
            print(" Wheel Type: "+row[1])
        if row[0]=="FEATURE_parkAssistEnabled":
            print(" Park Assist: "+row[1])
        if row[0]=="VAPI_hasFoldingMirrors":
            print(" Folding Mirrors: "+row[1])
        if row[0]=="VAPI_noKeylessEntry":
            if row[1]=="true":
                print(" Keyless Entry: false")
                print()
            else:
                print(" Keyless Entry: true")
                print()
        if row[0]=="VAPI_tpmsType":
            print(" TPMS Type: "+row[1])
        if row[0]=="VAPI_autopilot":
            print(" Autopilot: "+row[1])
        if row[0]=="CONN_cellIMEI":
            print(" IMEI Number: "+row[1])
        if row[0]=="CONN_cellConnected":
            print(" Cell Connection: "+row[1])
        if row[0]=="CONN_connectedToInternet":
            print(" Internet Connection: "+row[1])
        if row[0]=="CONN_vpnConnected":
            print(" Tesla Connection: "+row[1])
        if row[0]=="VAPI_performanceAddOn":
            print(" Performance AddOn: "+row[1])
    print()

def fillGeneralListBox():
    global rows
    if not isConnected:
        generalListbox.insert(END, "Bağlı değil. Özellikleri görmek için bağlanın.")
        return
    generalListbox.insert(END, " VIN Number: "+vin)
    for row in rows:
        row = row.split(",")
        if row[0]=="VAPI_chargerType":
            generalListbox.insert(END, " AC Charger Type: "+row[1])
        if row[0]=="VAPI_airSuspension":
            generalListbox.insert(END, " Air Suspension: "+row[1])
        if row[0]=="VAPI_frontFogLights":
            generalListbox.insert(END, " Front Fog Lights: "+row[1])
        if row[0]=="VAPI_rearFogLights":
            generalListbox.insert(END, " Rear Fog Lights: "+row[1])
        if row[0]=="VAPI_hasHomelink":
            generalListbox.insert(END, " Homelink: "+row[1])
        if row[0]=="VAPI_hasSunroof":
            generalListbox.insert(END, " Sunroof: "+row[1])
        if row[0]=="VAPI_hasPowerLiftgate":
            generalListbox.insert(END, " Power Liftgate: "+row[1])
        if row[0]=="FEATURE_blindspotWarningEnabled":
            generalListbox.insert(END, " Blindspot Warning: "+row[1])
        if row[0]=="VAPI_hasMemorySeats":
            generalListbox.insert(END, " Memory Seats: "+row[1])
        if row[0]=="VAPI_hasMemoryMirrors":
            generalListbox.insert(END, " Memory Mirrors: "+row[1])
        if row[0]=="SYS_IC_cpuHardware":
            generalListbox.insert(END, " CPU Hardware: "+row[1])
        if row[0]=="VAPI_hasSeatHeaters":
            generalListbox.insert(END, " Front Seat Heater: "+row[1])
        if row[0]=="VAPI_rearSeatHeaters":
            generalListbox.insert(END, " Rear Seat Heater: "+row[1])
        if row[0]=="VAPI_steeringWheelHeater":
            generalListbox.insert(END, " Steering Wheel Heater: "+row[1])
        if row[0]=="VAPI_fourWheelDrive":
            generalListbox.insert(END, " Four Wheel Drive: "+row[1])
        if row[0]=="VAPI_wheelType":
            generalListbox.insert(END, " Wheel Type: "+row[1])
        if row[0]=="FEATURE_parkAssistEnabled":
            generalListbox.insert(END, " Park Assist: "+row[1])
        if row[0]=="VAPI_hasFoldingMirrors":
            generalListbox.insert(END, " Folding Mirrors: "+row[1])
        if row[0]=="VAPI_noKeylessEntry":
            if row[1]=="true":
                generalListbox.insert(END, " Keyless Entry: false")
                generalListbox.insert(END, )
            else:
                generalListbox.insert(END, " Keyless Entry: true")
                generalListbox.insert(END, )
        if row[0]=="VAPI_tpmsType":
            generalListbox.insert(END, " TPMS Type: "+row[1])
        if row[0]=="VAPI_autopilot":
            generalListbox.insert(END, " Autopilot: "+row[1])
        if row[0]=="CONN_cellIMEI":
            generalListbox.insert(END, " IMEI Number: "+row[1])
        if row[0]=="CONN_cellConnected":
            generalListbox.insert(END, " Cell Connection: "+row[1])
        if row[0]=="CONN_connectedToInternet":
            generalListbox.insert(END, " Internet Connection: "+row[1])
        if row[0]=="CONN_vpnConnected":
            generalListbox.insert(END, " Tesla Connection: "+row[1])
        if row[0]=="VAPI_performanceAddOn":
            generalListbox.insert(END, " Performance AddOn: "+row[1])
    
def batterySoH(usableCapacity, effectiveCapacity):
    global rows
    print()
    for row in rows:
        row = row.split(",")
        if row[0]=="BMS_nominalFullPackEnergyRemaining":
            nominalFullPack = row[1]
            nominalFullPack = float(nominalFullPack[:len(nominalFullPack)-1])
            sohEffective = (nominalFullPack/effectiveCapacity)*100
            sohUsable = (nominalFullPack/usableCapacity)*100
            sohE = str('%.2f' %sohEffective)
            sohU = str('%.2f' %sohUsable)
            print(" Effective Battery SoH: %"+sohE)
            print(" Usable Battery SoH: %"+sohU)
    print()
    for row in rows:
        row = row.split(",")
        if row[0]=="VAPI_brickVoltageMax":
            brickMax = row[1]
            brickMax = float(brickMax[:len(brickMax)-1])
            bMax = str('%.2f' %brickMax)
            print(" Maximum Brick Voltage: "+bMax+"V")
        if row[0]=="VAPI_brickVoltageMin":
            brickMin = row[1]
            brickMin = float(brickMin[:len(brickMin)-1])
            bMin = str('%.2f' %brickMin)
            print(" Minimum Brick Voltage: "+bMin+"V")
    
    print()
    
    brickDelta = brickMax - brickMin
    brickError = brickDelta / brickMax
    
    bDelta = str('%.2f' %brickDelta)
    bError = str('%.2f' %brickError)
    
    print(" Maximum Potantial Difference: "+bDelta+"V")
    print(" Maximum Error: %"+bError)
    
    print()
    
    for row in rows:
        row = row.split(",")
        if row[0]=="VAPI_acChargerKwhTotal":
            acChargeCount = row[1]
            acChargeCount = float(acChargeCount[:len(acChargeCount)-1])
            acChargeCount = str('%.2f' %acChargeCount)
            print(" AC Charge Count: "+acChargeCount+" KWh")
        if row[0]=="VAPI_dcChargerKwhTotal":
            dcChargeCount = row[1]
            dcChargeCount = float(dcChargeCount[:len(dcChargeCount)-1])
            dcChargeCount = str('%.2f' %dcChargeCount)
            print(" DC Charge Count: "+dcChargeCount+" KWh")
        if row[0]=="VAPI_kWhDischargeCounter":
            kwDischargeCount = row[1]
            kwDischargeCount = float(kwDischargeCount[:len(kwDischargeCount)-1])
            kwDischargeCount = str('%.2f' %kwDischargeCount)
            print(" KWh Discharge Count: "+kwDischargeCount+" KWh")
    
    print()
    
    
def recentAlerts():
    global rows
    print()
    for row in rows:
        row = row.split(",")
        if row[0]=="carserver/recentAlerts":
            print(" 1 - Show Active Alerts Only")
            print(" 2 - Last 5 alerts")
            print(" 3 - Last 10 alerts")
            print(" 4 - Last 50 alerts")
            print(" 5 - Last 100 alerts")
            print(" 0 - Main Menu")
            print()
            alertsSelectionString = input("Select a number: ")
            alertNumber = 0
            activeOnly = False
            try:
                alertsSelection = int(alertsSelectionString)
                if alertsSelection == 1:
                    alertNumber = 100
                    activeOnly  = True
                elif alertsSelection == 2:
                    alertNumber = 5
                elif alertsSelection == 3:
                    alertNumber = 10
                elif alertsSelection == 4:
                    alertNumber = 50
                elif alertsSelection == 5:
                    alertNumber = 100
                elif alertsSelection == 0:
                    mainMenu()
                else:
                    print()
                    print("Invalid operation. Please enter a number between 0-5")
                    recentAlerts()
            except ValueError:
                print()
                print("Invalid operation. Please enter a number between 0-5")
                recentAlerts()

            if alertNumber != 0:
                if os.path.isfile(vin+"_alerts.txt"):
                    os.remove(vin+"_alerts.txt")
                text_file = open(vin+"_alerts.txt", "a")
                text_file.write("   ______          __      ____                        __ ")
                text_file.write("\n")
                text_file.write("  /_  __/__  _____/ /___ _/ __ \___  ____  ____  _____/ /_")
                text_file.write("\n")
                text_file.write("   / / / _ \/ ___/ / __ `/ /_/ / _ \/ __ \/ __ \/ ___/ __/")
                text_file.write("\n")
                text_file.write("  / / /  __(__  ) / /_/ / _, _/  __/ /_/ / /_/ / /  / /_  ")
                text_file.write("\n")
                text_file.write(" /_/  \___/____/_/\__,_/_/ |_|\___/ .___/\____/_/   \__/  ")
                text_file.write("\n")
                text_file.write("                                 /_/                      ")
                text_file.write("\n")
                text_file.write(" By ST-R Motors")
                text_file.write("\n")
                text_file.write("Service Alerts")
                text_file.write("\n")
                
                for i in range(alertNumber):
                    i = i+1
                    alertString = row[i]
                    alertString = alertString.split("@")
                    alertCode = alertString[0]
                    alertTime = alertString[1]
                    alertDesc = alertString[2]
                    alertEnd = alertString[3]
                    
                    alertTime = alertTime.split("T")
                    if alertEnd != "":
                        alertEnd = alertEnd.split("T")
                    
                    alertLogs = []
                    
                    if activeOnly == False:
                        print()
                        alertLogs.append("\n")
                        print("[Alert "+str(i)+"]")
                        alertLogs.append("[Arıza "+str(i)+"]")
                        alertLogs.append("\n")
                        print()
                        alertLogs.append("\n")
                        print(" Alert Code: " + alertCode)
                        alertLogs.append("Arıza Kodu: " + alertCode)
                        alertLogs.append("\n")
                        print(" Alert Description: " + alertDesc)
                        alertLogs.append("Arıza Açıklaması: " + alertDesc)
                        alertLogs.append("\n")
                        print(" Alert Date: " + alertTime[0])
                        alertLogs.append("Arıza Tarihi: " + alertTime[0])
                        alertLogs.append("\n")
                        print(" Alert Time: " + alertTime[1])
                        alertLogs.append("Arıza Saati: " + alertTime[1])
                        alertLogs.append("\n")
                        if alertEnd != "":
                            print(" Alert End Date: " + alertEnd[0])
                            alertLogs.append("Arıza Tarihi: " + alertTime[0])
                            alertLogs.append("\n")
                            print(" Alert End Time: " + alertEnd[1])
                            alertLogs.append("Arıza Saati: " + alertTime[1])
                            alertLogs.append("\n")
                        else:
                            print(" [!] Alert is active!")
                            #alertLogs.append(" [!] Arıza Aktif!")
                            alertLogs.append("\n")
                        
                        for alert in alertLogs:
                                       text_file.write(alert)
                    else:
                        if alertEnd == "":
                            print()
                            alertLogs.append("\n")
                            print(" Alert Code: " + alertCode)
                            alertLogs.append("Arıza Kodu: " + alertCode)
                            alertLogs.append("\n")
                            print(" Alert Description: " + alertDesc)
                            alertLogs.append("Arıza Açıklaması: " + alertDesc)
                            alertLogs.append("\n")
                            print(" Alert Date: " + alertTime[0])
                            alertLogs.append("Arıza Tarihi: " + alertTime[0])
                            alertLogs.append("\n")
                            print(" Alert Time: " + alertTime[1])
                            alertLogs.append("Arıza Saati: " + alertTime[1])
                            alertLogs.append("\n")
                            print(" [!] Alert is active!")
                            #alertLogs.append(" [!] Arıza Aktif!")
                            alertLogs.append("\n")
                            for alert in alertLogs:
                                           text_file.write(alert)
                
                print()
                print(" Output file saved as "+vin+"_alerts.txt")
                text_file.close()
    print()

def fillWarningListBox():
    global rows
    global numberOfAlerts
    if not isConnected:
        warningListbox.insert(END, "Bağlı değil. Özellikleri görmek için bağlanın.")
        return
    
    for row in rows:
        row = row.split(",")
        if row[0]=="carserver/recentAlerts":
            alertNumber = 100
            activeOnly = True

            try:
                for i in range(alertNumber):
                    i = i+1
                    alertString = row[i]
                    alertString = alertString.split("@")
                    alertCode = alertString[0]
                    alertTime = alertString[1]
                    alertDesc = alertString[2]
                    alertEnd = alertString[3]
                    
                    alertTime = alertTime.split("T")
                    if alertEnd != "":
                        alertEnd = alertEnd.split("T")
                    
                    alertLogs = []
                    
                    if alertEnd == "":
                        warningListbox.insert(END, alertCode)
                        warningListbox.insert(END, "Açıklama: " + alertDesc)
                        warningListbox.insert(END, "Tarih: " + alertTime[0])
                        warningListbox.insert(END, "Saat: " + alertTime[1])
                        warningListbox.insert(END, "")
                        numberOfAlerts = numberOfAlerts+1
            except:
                cprint("Input/Output hatası. Dosyaları temizleyip tekrar deneyin.")
    
def exportReport():
    global rows
    global diagVitals
    global usableCapacity
    global effectiveCapacity
    
    if isConnected:
        #GENERAL INFO
        
        cprint("(1/3) Genel bilgiler toplanıyor...")
        
        for row in rows:
            row = row.split(",")
            if row[0]=="VAPI_chargerType":
                writeToReport("$actype",row[1])
            if row[0]=="VAPI_airSuspension":
                writeToReport("$suspension",row[1])
            if row[0]=="VAPI_frontFogLights":
                writeToReport("$frontfogs",row[1])
            if row[0]=="VAPI_rearFogLights":
                writeToReport("$rearfogs",row[1])
            if row[0]=="VAPI_hasHomelink":
                writeToReport("$homelink",row[1])
            if row[0]=="VAPI_hasSunroof":
                writeToReport("$sunroof",row[1])
            if row[0]=="VAPI_hasPowerLiftgate":
                writeToReport("$powerliftgate",row[1])
            if row[0]=="FEATURE_blindspotWarningEnabled":
                writeToReport("$blindspot",row[1])
            if row[0]=="VAPI_hasMemorySeats":
                writeToReport("$memoryseats",row[1])
            if row[0]=="VAPI_hasMemoryMirrors":
                writeToReport("$memorymirrors",row[1])
            if row[0]=="SYS_IC_cpuHardware":
                writeToReport("$cpuhw",row[1])
            if row[0]=="VAPI_hasSeatHeaters":
                writeToReport("$frontseatheat",row[1])
            if row[0]=="VAPI_rearSeatHeaters":
                writeToReport("$rearseatheat",row[1])
            if row[0]=="VAPI_steeringWheelHeater":
                writeToReport("$steeringheat",row[1])
            if row[0]=="VAPI_fourWheelDrive":
                writeToReport("$isawd",row[1])
            if row[0]=="VAPI_wheelType":
                writeToReport("$wheeltype",row[1])
            if row[0]=="FEATURE_parkAssistEnabled":
                writeToReport("$parkassist",row[1])
            if row[0]=="VAPI_hasFoldingMirrors":
                writeToReport("$foldingmirrors",row[1])
            if row[0]=="VAPI_noKeylessEntry":
                if row[1]=="true":
                    writeToReport("$keylessentry","false")
                else:
                    writeToReport("$keylessentry","true")
            if row[0]=="VAPI_tpmsType":
                writeToReport("$tpmstype",row[1])
            if row[0]=="VAPI_autopilot":
                writeToReport("$autopilot",row[1])
            if row[0]=="CONN_cellIMEI":
                writeToReport("$imeinumber",row[1])
            if row[0]=="CONN_connectedToInternet":
                writeToReport("$internetconnection",row[1])
            if row[0]=="CONN_vpnConnected":
                writeToReport("$teslaconnection",row[1])
            if row[0]=="VAPI_performanceAddOn":
                writeToReport("$performanceaddon",row[1])
        writeToReport("true","Mevcut")
        writeToReport("false","Mevcut Değil")
        writeToReport("--","")
        writeToReport("None","Mevcut Değil")
        
        #BATTERY INFO
        
        cprint("(2/3) Batarya verileri toplanıyor...")
        
        for row in rows:
            row = row.split(",")
            if row[0]=="BMS_nominalFullPackEnergyRemaining":
                nominalFullPack = row[1]
                nominalFullPack = float(nominalFullPack[:len(nominalFullPack)-1])
                sohEffective = (nominalFullPack/effectiveCapacity)*100
                sohUsable = (nominalFullPack/usableCapacity)*100
                sohE = str('%.2f' %sohEffective)
                sohU = str('%.2f' %sohUsable)
                writeToReport("$effectivesoh","%"+sohE)
                writeToReport("$usablesoh","%"+sohU)
        for row in rows:
            row = row.split(",")
            if row[0]=="VAPI_brickVoltageMax":
                brickMax = row[1]
                brickMax = float(brickMax[:len(brickMax)-1])
                bMax = str('%.2f' %brickMax)
                writeToReport("$maxbrick",bMax+"V")
            if row[0]=="VAPI_brickVoltageMin":
                brickMin = row[1]
                brickMin = float(brickMin[:len(brickMin)-1])
                bMin = str('%.2f' %brickMin)
                writeToReport("$minbrick",bMin+"V")
        
        brickDelta = brickMax - brickMin
        brickError = brickDelta / brickMax
        
        bDelta = str('%.2f' %brickDelta)
        bError = str('%.2f' %brickError)
        
        writeToReport("$maxpotdiff",bDelta+"V")
        writeToReport("$maxerror","%"+bError)
            
        for row in rows:
            row = row.split(",")
            if row[0]=="VAPI_acChargerKwhTotal":
                acChargeCount = row[1]
                acChargeCount = float(acChargeCount[:len(acChargeCount)-1])
                acChargeCount = str('%.2f' %acChargeCount)
                writeToReport("$ACchargecount",acChargeCount+" KWh")
            if row[0]=="VAPI_dcChargerKwhTotal":
                dcChargeCount = row[1]
                dcChargeCount = float(dcChargeCount[:len(dcChargeCount)-1])
                dcChargeCount = str('%.2f' %dcChargeCount)
                writeToReport("$DCchargecount",dcChargeCount+" KWh")
            if row[0]=="VAPI_kWhDischargeCounter":
                kwDischargeCount = row[1]
                kwDischargeCount = float(kwDischargeCount[:len(kwDischargeCount)-1])
                kwDischargeCount = str('%.2f' %kwDischargeCount)
                writeToReport("$dischargecount",kwDischargeCount+" KWh")
        
        #ALERT INFO
        
        cprint("(3/3) Servis arızaları toplanıyor...")
        
        for row in rows:
            row = row.split(",")
            if row[0]=="carserver/recentAlerts":
                alertNumber = 100
                activeOnly = True

                try:
                    if os.path.isfile("alerts_cache"):
                        os.remove("alerts_cache")
                    text_file = open("alerts_cache", "a")
                    
                    for i in range(alertNumber):
                        i = i+1
                        alertString = row[i]
                        alertString = alertString.split("@")
                        alertCode = alertString[0]
                        alertTime = alertString[1]
                        alertDesc = alertString[2]
                        alertEnd = alertString[3]
                        
                        alertTime = alertTime.split("T")
                        if alertEnd != "":
                            alertEnd = alertEnd.split("T")
                        
                        alertLogs = []
                        
                        if alertEnd == "":
                            alertLogs.append("\n")
                            alertLogs.append("Arıza Kodu: " + alertCode)
                            alertLogs.append("\n")
                            
                            alertLogs.append("Arıza Açıklaması: " + alertDesc)
                            alertLogs.append("\n")
                            
                            alertLogs.append("Arıza Tarihi: " + alertTime[0])
                            alertLogs.append("\n")
                            
                            alertLogs.append("Arıza Saati: " + alertTime[1])
                            alertLogs.append("\n")
                            alertLogs.append("\n")
                            
                            for alert in alertLogs:
                                text_file.write(alert)
                    
                    text_file.close()
                    with open("alerts_cache", 'r') as f:
                        writeToReport("$activeservicealerts",f.read())
                    
                    os.remove("alerts_cache")
                except:
                    cprint("Input/Output error. Try again after cleaning files.")
        
        print()
        cprint("All information is collected.")
        
        global vin
        
        saveReportName = vin + "_Report.docx"
        shutil.copy('report_cache', saveReportName)
        
        print()
        cprint("Rapor ./" + saveReportName+ " ismiyle kaydedildi.")
        os.remove('report_cache')
        
    else:
        cprint("Lütfen bağlantı kurup tekrar deneyin.")

def revealPins():
    global rows
    print()
    for row in rows:
        row = row.split(",")
        if row[0]=="GUI_PINToDrivePassword":
            ptd = row[1]
            print(" PIN to Drive Password: " + ptd)
    for row in rows:
        row = row.split(",")
        if row[0]=="GUI_gloveboxPassword":
            gb = row[1]
            print(" Glovebox Password: " + gb)
    for row in rows:
        row = row.split(",")
        if row[0]=="GUI_speedLimitModePassword":
            sl = row[1]
            print(" Speed Limit Password: " + sl)
    for row in rows:
        row = row.split(",")
        if row[0]=="GUI_valetModePassword":
            vm = row[1]
            print(" Valet Mode Password: " + vm)

def revealSpotify():
    global rows
    print()
    for row in rows:
        row = row.split(",")
        if row[0]=="MEDIA_spotifyUsername":
            username = row[1]
            print(" Spotify Username: " + username)
    for row in rows:
        row = row.split(",")
        if row[0]=="MEDIA_spotifyPassword":
            password = row[1]
            print(" Spotify Password: " + password)

def revealWifi():
    global rows
    print()
    for row in rows:
        if "LINK_wifiKnownNetworks," in row:
            row = row.split(",[ { ")
            row = row[1]
            row = row.split("}, {")
            for network in row:
                ssid = ""
                spass = ""
                network = network.split(",")
                for netData in network:
                    if "ssid" in netData:
                        netData = netData.split(":")
                        netData = netData[1].split("\"")
                        ssid = netData[1]
                
                for netData in network:
                    if "key" in netData:
                        netData = netData.split(":")
                        netData = netData[1].split("\"")
                        spass = netData[1]
                
                print("SSID: "+ssid)
                print("Password: "+spass)
                print()

activePage = "general"

OUTPUT_PATH = Path(__file__).parent
ASSETS_PATH = OUTPUT_PATH / Path(r"assets")

def relative_to_assets(path: str) -> Path:
    return ASSETS_PATH / Path(path)
    
global window
window = Tk()
window.iconbitmap(relative_to_assets("icon.ico"))
window.title('TeslaReport')


def goBatteryPage():
    global rows
    global activePage
    if activePage == "general":
        canvas.delete(modelNameLabel)
        generalListbox.destroy()
        generalScrollbar.destroy()
        generalButton.config(image=generalUnclickedBG)
        batteryButton.config(image=batteryClickedBG)
        
    elif activePage == "warning":
        canvas.delete(alertNumberLabel)
        warningButton.config(image=warningUnclickedBG)
        batteryButton.config(image=batteryClickedBG)
        warningListbox.destroy()
        warningScrollbar.destroy()
        
    else:
        return
    
    if activePage != "battery":
        canvas.itemconfig(bodyBG, image=batteryBodyBG)
        canvas.itemconfig(topBarImage, image=batteryTopBG)
        
        if isConnected:
            for row in rows:
                row = row.split(",")
                if row[0]=="VAPI_brickVoltageMax":
                    brickMax = row[1]
                    brickMax = float(brickMax[:len(brickMax)-1])
                    bMax = str('%.2f' %brickMax)
                if row[0]=="VAPI_brickVoltageMin":
                    brickMin = row[1]
                    brickMin = float(brickMin[:len(brickMin)-1])
                    bMin = str('%.2f' %brickMin)
                    
                if row[0]=="VAPI_acChargerKwhTotal":
                    acChargeCount = row[1]
                    acChargeCount = float(acChargeCount[:len(acChargeCount)-1])
                    acChargeCount = str('%.2f' %acChargeCount)
                if row[0]=="VAPI_dcChargerKwhTotal":
                    dcChargeCount = row[1]
                    dcChargeCount = float(dcChargeCount[:len(dcChargeCount)-1])
                    dcChargeCount = str('%.2f' %dcChargeCount)
                if row[0]=="VAPI_kWhDischargeCounter":
                    kwDischargeCount = row[1]
                    kwDischargeCount = float(kwDischargeCount[:len(kwDischargeCount)-1])
                    kwDischargeCount = str('%.2f' %kwDischargeCount)
            
            brickDelta = brickMax - brickMin
            brickError = brickDelta / brickMax
            
            bDelta = str('%.2f' %brickDelta)
            bError = str('%.2f' %brickError)

            global usableBatteryHealthLabel
            usableBatteryHealthLabel = canvas.create_text(
                420.0,
                138.0,
                anchor="nw",
                text="%"+str(usableCapacity),
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )
            
            global effectiveBatteryHealthLabel
            effectiveBatteryHealthLabel = canvas.create_text(
                420.0,
                116.0,
                anchor="nw",
                text="%"+str(effectiveCapacity),
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )

            global minBrickVoltageLabel
            minBrickVoltageLabel = canvas.create_text(
                420.0,
                251.0,
                anchor="nw",
                text=bMin+"V",
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )

            global maxBrickVoltageLabel
            maxBrickVoltageLabel = canvas.create_text(
                420.0,
                229.0,
                anchor="nw",
                text=bMax+"V",
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )

            global maxErrorLabel
            maxErrorLabel = canvas.create_text(
                420.0,
                364.0,
                anchor="nw",
                text="%"+bError,
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )
            
            global maxPotDiffLabel
            maxPotDiffLabel = canvas.create_text(
                420.0,
                342.0,
                anchor="nw",
                text=bDelta+"V",
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )

            global acCounterLabel
            acCounterLabel = canvas.create_text(
                326.0,
                478.0,
                anchor="nw",
                text=acChargeCount+" KWh",
                fill="#000000",
                font=("Inter Bold", 13 * -1)
            )

            global dcCounterLabel
            dcCounterLabel = canvas.create_text(
                517.0,
                478.0,
                anchor="nw",
                text=dcChargeCount+" KWh",
                fill="#000000",
                font=("Inter Bold", 13 * -1)
            )

            global dischargeCounterLabel
            dischargeCounterLabel = canvas.create_text(
                708.0,
                478.0,
                anchor="nw",
                text=kwDischargeCount+" KWh",
                fill="#000000",
                font=("Inter Bold", 13 * -1)
            )
        else:
            usableBatteryHealthLabel = canvas.create_text(
                420.0,
                138.0,
                anchor="nw",
                text="N/A",
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )
            
            effectiveBatteryHealthLabel = canvas.create_text(
                420.0,
                116.0,
                anchor="nw",
                text="N/A",
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )

            minBrickVoltageLabel = canvas.create_text(
                420.0,
                251.0,
                anchor="nw",
                text="N/A",
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )

            maxBrickVoltageLabel = canvas.create_text(
                420.0,
                229.0,
                anchor="nw",
                text="N/A",
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )

            maxErrorLabel = canvas.create_text(
                420.0,
                364.0,
                anchor="nw",
                text="N/A",
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )
            
            maxPotDiffLabel = canvas.create_text(
                420.0,
                342.0,
                anchor="nw",
                text="N/A",
                fill="#000000",
                font=("Inter Bold", 11 * -1)
            )

            acCounterLabel = canvas.create_text(
                326.0,
                478.0,
                anchor="nw",
                text="N/A",
                fill="#000000",
                font=("Inter Bold", 13 * -1)
            )

            dcCounterLabel = canvas.create_text(
                517.0,
                478.0,
                anchor="nw",
                text="N/A",
                fill="#000000",
                font=("Inter Bold", 13 * -1)
            )

            dischargeCounterLabel = canvas.create_text(
                708.0,
                478.0,
                anchor="nw",
                text="N/A",
                fill="#000000",
                font=("Inter Bold", 13 * -1)
            )
        activePage = "battery"
    
def goGeneralPage():
    global rows
    global activePage
    if activePage == "battery":
        canvas.delete(usableBatteryHealthLabel)
        canvas.delete(effectiveBatteryHealthLabel)
        canvas.delete(minBrickVoltageLabel)
        canvas.delete(maxBrickVoltageLabel)
        canvas.delete(maxPotDiffLabel)
        canvas.delete(maxErrorLabel)
        canvas.delete(acCounterLabel)
        canvas.delete(dcCounterLabel)
        canvas.delete(dischargeCounterLabel)
    
        batteryButton.config(image=batteryUnclickedBG)
        generalButton.config(image=generalClickedBG)
        
    elif activePage == "warning":
        canvas.delete(alertNumberLabel)
        warningButton.config(image=warningUnclickedBG)
        generalButton.config(image=generalClickedBG)
        warningListbox.destroy()
        warningScrollbar.destroy()
        
    else:
        return
        
    if activePage != "general":
        canvas.itemconfig(bodyBG, image=generalBodyBG)
        canvas.itemconfig(topBarImage, image=generalTopBG)
        
        global modelNameLabel
        if isConnected:
            modelNameLabel = canvas.create_text(
                685.0,
                213.0,
                anchor="nw",
                text=carType+trimName,
                fill="#000000",
                font=("Inter SemiBold", 11 * -1)
            )
        else:
            modelNameLabel = canvas.create_text(
                685.0,
                213.0,
                anchor="nw",
                text="Bağlı Değil!",
                fill="#000000",
                font=("Inter SemiBold", 11 * -1)
            )
        
        global generalScrollbar
        generalScrollbar = Scrollbar(window, orient= 'vertical')
        generalScrollbar.pack(side= RIGHT, fill= BOTH)

        global generalListbox
        generalListbox = Listbox(window, width= 350, height=200, borderwidth=0, highlightthickness=0, font= ("Inter SemiBold", 13 * -1))
        generalListbox.place(
            x=250,
            y=73,
            width=302,
            height=429
        )

        fillGeneralListBox() 

        generalListbox.config(yscrollcommand= generalScrollbar.set)

        generalScrollbar.config(command= generalListbox.yview)
        generalScrollbar.place(x=552, y=63, height=449, width=10)
        
        activePage = "general"

def goWarningPage():
    global rows
    global activePage
    if activePage == "battery":
        canvas.delete(usableBatteryHealthLabel)
        canvas.delete(effectiveBatteryHealthLabel)
        canvas.delete(minBrickVoltageLabel)
        canvas.delete(maxBrickVoltageLabel)
        canvas.delete(maxPotDiffLabel)
        canvas.delete(maxErrorLabel)
        canvas.delete(acCounterLabel)
        canvas.delete(dcCounterLabel)
        canvas.delete(dischargeCounterLabel)
    
        batteryButton.config(image=batteryUnclickedBG)
        warningButton.config(image=warningClickedBG)
    
    elif activePage == "general":
        canvas.delete(modelNameLabel)
        generalListbox.destroy()
        generalScrollbar.destroy()
        
        generalButton.config(image=generalUnclickedBG)
        warningButton.config(image=warningClickedBG)
    
    else:
        return
    
    if activePage != "warning":
        canvas.itemconfig(bodyBG, image=warningBodyBG)
        canvas.itemconfig(topBarImage, image=warningTopBG)
        
        global warningScrollbar
        warningScrollbar = Scrollbar(window, orient= 'vertical')
        warningScrollbar.pack(side= RIGHT, fill= BOTH)

        global warningListbox
        warningListbox = Listbox(window, width= 350, height=200, borderwidth=0, highlightthickness=0, font= ("Inter SemiBold", 13 * -1))
        warningListbox.place(
            x=250,
            y=73,
            width=302,
            height=429
        )

        global numberOfAlerts
        numberOfAlerts = 0
        
        fillWarningListBox() 

        warningListbox.config(yscrollcommand= warningScrollbar.set)

        warningScrollbar.config(command= warningListbox.yview)
        warningScrollbar.place(x=552, y=63, height=449, width=10)
        
        global alertNumberLabel
        if isConnected:
            alertNumberLabel = canvas.create_text(
                670.0,
                298.0,
                anchor="nw",
                text=str(numberOfAlerts)+" Arıza Aktif",
                fill="#363636",
                font=("Inter SemiBold", 24 * -1)
            )
        else:
            alertNumberLabel = canvas.create_text(
                670.0,
                298.0,
                anchor="nw",
                text="Bağlı Değil!",
                fill="#363636",
                font=("Inter SemiBold", 24 * -1)
            )
        
        activePage = "warning"

window.geometry("897x555")
window.configure(bg = "#FFFFFF")

canvas = Canvas(
    window,
    bg = "#FFFFFF",
    height = 555,
    width = 897,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
)

canvas.place(x = 0, y = 0)
canvas.create_rectangle(
    0.0,
    54.0,
    229.0,
    553.0,
    fill="#141414",
    outline="")

batteryBodyBG = PhotoImage(
    file=relative_to_assets("batterybody.png"))
generalBodyBG = PhotoImage(
    file=relative_to_assets("generalbody.png"))
warningBodyBG = PhotoImage(
    file=relative_to_assets("warningbody.png"))
    
bodyBG = canvas.create_image(
    563.0,
    289.0,
    image=generalBodyBG
)

notConnectedBG = PhotoImage(
    file=relative_to_assets("notconnected.png"))
connectedBG = PhotoImage(
    file=relative_to_assets("connected.png"))
    
connectionButton = Button(
    image=notConnectedBG,
    borderwidth=0,
    highlightthickness=0,
    command=lambda: threading.Thread(target=tryConnection).start(),
    relief="flat",
    compound="center"
)
connectionButton.place(
    x=0.0,
    y=525.0,
    width=229.0,
    height=30.0
)

reportButtonBG = PhotoImage(
    file=relative_to_assets("report.png"))
reportButton = Button(
    image=reportButtonBG,
    borderwidth=0,
    highlightthickness=0,
    command=lambda: threading.Thread(target=exportReport).start(),
    relief="flat"
)
reportButton.place(
    x=0.0,
    y=459.0,
    width=229.0,
    height=66.0
)

warningUnclickedBG = PhotoImage(
    file=relative_to_assets("warning.png"))
warningClickedBG = PhotoImage(
    file=relative_to_assets("warningclicked.png"))
    
warningButton = Button(
    image=warningUnclickedBG,
    borderwidth=0,
    highlightthickness=0,
    command=lambda: goWarningPage(),
    relief="flat"
)
warningButton.place(
    x=0.0,
    y=186.0,
    width=229.0,
    height=66.0
)

batteryClickedBG = PhotoImage(
    file=relative_to_assets("batteryclicked.png"))
batteryUnclickedBG = PhotoImage(
    file=relative_to_assets("battery.png"))
    
batteryButton = Button(
    image=batteryUnclickedBG,
    borderwidth=0,
    highlightthickness=0,
    command=lambda: goBatteryPage(),
    relief="flat"
)
batteryButton.place(
    x=0.0,
    y=120.0,
    width=229.0,
    height=66.0
)

generalClickedBG = PhotoImage(
    file=relative_to_assets("generalclicked.png"))
generalUnclickedBG = PhotoImage(
    file=relative_to_assets("general.png"))
    
generalButton = Button(
    image=generalClickedBG,
    borderwidth=0,
    highlightthickness=0,
    command=lambda: goGeneralPage(),
    relief="flat"
)
generalButton.place(
    x=0.0,
    y=54.0,
    width=229.0,
    height=66.0
)

generalTopBG = PhotoImage(
    file=relative_to_assets("generaltop.png"))
batteryTopBG = PhotoImage(
    file=relative_to_assets("batterytop.png"))
warningTopBG = PhotoImage(
    file=relative_to_assets("warningtop.png"))
    
topBarImage = canvas.create_image(
    448.0,
    27.0,
    image=generalTopBG
)

bottomBarBG = PhotoImage(
    file=relative_to_assets("bottom.png"))
bottomBar = canvas.create_image(
    563.0,
    540.0,
    image=bottomBarBG
)

modelNameLabel = canvas.create_text(
    705.0,
    213.0,
    anchor="nw",
    text="Bağlı Değil!",
    fill="#000000",
    font=("Inter SemiBold", 11 * -1)
)

bottomConsole = canvas.create_text(
    242.0,
    533.0,
    anchor="nw",
    text="Bağlı değil.",
    fill="#FFFFFF",
    font=("Inter Regular", 13 * -1)
)

generalScrollbar = Scrollbar(window, orient= 'vertical')
generalScrollbar.pack(side= RIGHT, fill= BOTH)

generalListbox = Listbox(window, borderwidth=0, highlightthickness=0, font= ("Inter SemiBold", 13 * -1))
generalListbox.place(
    x=250,
    y=73,
    width=302,
    height=429
)

fillGeneralListBox() 

generalListbox.config(yscrollcommand= generalScrollbar.set)

generalScrollbar.config(command= generalListbox.yview)
generalScrollbar.place(x=552, y=63, height=449, width=10)

window.resizable(False, False)
window.mainloop()